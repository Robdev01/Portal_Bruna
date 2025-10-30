from docx import Document
import re

def preencher_modelo(modelo_path: str, dados: dict, output_path: str):
    """
    Substitui {{CHAVE}} preservando pontuação e formatação.
    Se o placeholder atravessa vários runs, o valor é escrito INTEIRO no primeiro run,
    garantindo que todo o valor herde o mesmo estilo (ex.: todo em negrito).
    """

    doc = Document(modelo_path)
    dados_str = {k: str(v or "") for k, v in dados.items()}

    # casa {{CHAVE}} com tolerância a espaços internos e pontuação logo após
    def pattern_for(key: str):
        # {{   CHAVE   }} seguido de , . ; : ) espaço ou fim da linha
        return re.compile(rf"\{{\{{\s*{re.escape(key)}\s*\}}\}}(?=[,.;:\s)]|$)")

    def replace_in_paragraph(p):
        if not p.runs:
            return

        # Texto combinado e mapa de posições -> (idx_run, idx_char_no_run)
        full = []
        idxmap = []  # para cada posição em full, guarda (ri, ci)
        for ri, run in enumerate(p.runs):
            for ci, ch in enumerate(run.text):
                full.append(ch)
                idxmap.append((ri, ci))
        full = "".join(full)

        # Executa substituições para todas as chaves
        for key, value in dados_str.items():
            pat = pattern_for(key)

            # Iteramos por matches sucessivos; após cada replace, recalculamos 'full' e 'idxmap'
            while True:
                m = pat.search(full)
                if not m:
                    break

                a, b = m.start(), m.end()  # intervalo do placeholder em 'full'
                (ri_start, ci_start) = idxmap[a]
                (ri_end,   ci_end)   = idxmap[b-1]  # último char do placeholder

                # Runs que participam do placeholder
                first_run = p.runs[ri_start]
                last_run  = p.runs[ri_end]

                # Partes antes e depois, preservando pontuação no sufixo
                prefix = first_run.text[:ci_start]
                suffix = last_run.text[ci_end+1:] if (ri_end == ri_start) else last_run.text[ci_end+1:]

                # 1) escreve tudo no primeiro run: prefix + VALUE + suffix
                #    (o first_run mantém sua formatação: negrito, cor, tamanho etc.)
                first_run.text = prefix + value + (suffix if ri_end == ri_start else "")

                # 2) esvazia os runs intermediários
                for ri in range(ri_start+1, ri_end):
                    p.runs[ri].text = ""

                # 3) se o placeholder atravessava vários runs, limpa a parte antes/sufixo nos extremos
                if ri_end != ri_start:
                    # limpa tudo do last_run (ficará somente o sufixo fora do placeholder)
                    last_run.text = suffix

                # Recalcula 'full' e 'idxmap' após a alteração
                full = []
                idxmap = []
                for r_i, run in enumerate(p.runs):
                    for c_i, ch in enumerate(run.text):
                        full.append(ch)
                        idxmap.append((r_i, c_i))
                full = "".join(full)

        # fim for key

    def walk_tables(tbls):
        for t in tbls:
            for row in t.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para)
                    walk_tables(cell.tables)

    # Parágrafos de nível superior
    for para in doc.paragraphs:
        replace_in_paragraph(para)

    # Dentro de tabelas (recursivo)
    walk_tables(doc.tables)

    doc.save(output_path)
